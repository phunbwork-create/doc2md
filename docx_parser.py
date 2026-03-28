"""
Doc2MD Parser Engine
Phân tích sâu file .docx và xuất ra Markdown template chi tiết.
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import re


def emu_to_cm(emu_val):
    """Chuyển đổi EMU sang cm."""
    if emu_val is None:
        return None
    return round(emu_val / 914400 * 2.54, 2)


def emu_to_pt(emu_val):
    """Chuyển đổi EMU sang point."""
    if emu_val is None:
        return None
    return round(emu_val / 12700, 1)


def twips_to_cm(twips_val):
    """Chuyển đổi twips sang cm."""
    if twips_val is None:
        return None
    return round(twips_val / 567, 2)


def get_alignment_name(alignment):
    """Trả về tên alignment dạng text."""
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "Căn trái (Left)",
        WD_ALIGN_PARAGRAPH.CENTER: "Căn giữa (Center)",
        WD_ALIGN_PARAGRAPH.RIGHT: "Căn phải (Right)",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "Căn đều (Justify)",
        WD_ALIGN_PARAGRAPH.DISTRIBUTE: "Phân bố (Distribute)",
    }
    return mapping.get(alignment, "Mặc định (Left)")


def get_line_spacing_rule_name(rule):
    """Trả về tên line spacing rule."""
    mapping = {
        WD_LINE_SPACING.SINGLE: "Đơn (Single)",
        WD_LINE_SPACING.ONE_POINT_FIVE: "1.5 dòng",
        WD_LINE_SPACING.DOUBLE: "Đôi (Double)",
        WD_LINE_SPACING.AT_LEAST: "Tối thiểu (At Least)",
        WD_LINE_SPACING.EXACTLY: "Chính xác (Exactly)",
        WD_LINE_SPACING.MULTIPLE: "Bội số (Multiple)",
    }
    return mapping.get(rule, "Không xác định")


def parse_color(color_obj):
    """Trích xuất thông tin màu sắc."""
    if color_obj is None:
        return None
    try:
        if color_obj.rgb:
            return f"#{color_obj.rgb}"
    except (AttributeError, TypeError):
        pass
    try:
        if color_obj.theme_color:
            return f"Theme: {color_obj.theme_color}"
    except (AttributeError, TypeError):
        pass
    return None


def extract_run_formatting(run):
    """Trích xuất formatting chi tiết của một Run."""
    fmt = {}
    if run.font.name:
        fmt["font"] = run.font.name
    if run.font.size:
        fmt["size"] = f"{run.font.size.pt}pt"
    if run.bold:
        fmt["bold"] = True
    if run.italic:
        fmt["italic"] = True
    if run.underline:
        fmt["underline"] = True
    if run.font.strike:
        fmt["strikethrough"] = True
    if run.font.superscript:
        fmt["superscript"] = True
    if run.font.subscript:
        fmt["subscript"] = True

    color = parse_color(run.font.color)
    if color:
        fmt["color"] = color

    try:
        highlight = run.font.highlight_color
        if highlight:
            fmt["highlight"] = str(highlight)
    except Exception:
        pass

    return fmt


def format_run_info(fmt_dict):
    """Chuyển dict formatting thành chuỗi mô tả."""
    parts = []
    if "font" in fmt_dict:
        parts.append(f"Font: {fmt_dict['font']}")
    if "size" in fmt_dict:
        parts.append(f"Size: {fmt_dict['size']}")
    if fmt_dict.get("bold"):
        parts.append("**In đậm**")
    if fmt_dict.get("italic"):
        parts.append("*In nghiêng*")
    if fmt_dict.get("underline"):
        parts.append("Gạch chân")
    if fmt_dict.get("strikethrough"):
        parts.append("~~Gạch ngang~~")
    if fmt_dict.get("superscript"):
        parts.append("Chỉ số trên")
    if fmt_dict.get("subscript"):
        parts.append("Chỉ số dưới")
    if "color" in fmt_dict:
        parts.append(f"Màu: {fmt_dict['color']}")
    if "highlight" in fmt_dict:
        parts.append(f"Highlight: {fmt_dict['highlight']}")
    return ", ".join(parts) if parts else "Mặc định"


def extract_paragraph_formatting(paragraph):
    """Trích xuất formatting chi tiết của Paragraph."""
    pf = paragraph.paragraph_format
    info = {}

    info["alignment"] = get_alignment_name(pf.alignment)

    if pf.left_indent:
        info["left_indent"] = f"{emu_to_cm(pf.left_indent)} cm"
    if pf.right_indent:
        info["right_indent"] = f"{emu_to_cm(pf.right_indent)} cm"
    if pf.first_line_indent:
        info["first_line_indent"] = f"{emu_to_cm(pf.first_line_indent)} cm"

    if pf.space_before:
        info["space_before"] = f"{pf.space_before.pt} pt"
    if pf.space_after:
        info["space_after"] = f"{pf.space_after.pt} pt"

    if pf.line_spacing:
        if pf.line_spacing_rule:
            info["line_spacing"] = f"{pf.line_spacing} ({get_line_spacing_rule_name(pf.line_spacing_rule)})"
        else:
            info["line_spacing"] = str(pf.line_spacing)

    if pf.keep_together:
        info["keep_together"] = True
    if pf.keep_with_next:
        info["keep_with_next"] = True
    if pf.page_break_before:
        info["page_break_before"] = True
    if pf.widow_control:
        info["widow_control"] = True

    return info


def detect_list_info(paragraph):
    """Phát hiện thông tin danh sách (bullet/numbered)."""
    pPr = paragraph._element.find(qn('w:pPr'))
    if pPr is None:
        return None

    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        return None

    ilvl_elem = numPr.find(qn('w:ilvl'))
    numId_elem = numPr.find(qn('w:numId'))

    level = int(ilvl_elem.get(qn('w:val'))) if ilvl_elem is not None else 0
    numId = int(numId_elem.get(qn('w:val'))) if numId_elem is not None else 0

    return {
        "type": "numbered" if numId > 0 else "bullet",
        "level": level,
        "numId": numId
    }


def detect_images_in_paragraph(paragraph):
    """Phát hiện hình ảnh trong paragraph."""
    images = []
    for run in paragraph.runs:
        drawing_elems = run._element.findall(qn('w:drawing'))
        for drawing in drawing_elems:
            # Tìm extent để lấy kích thước
            extents = drawing.findall('.//' + qn('wp:extent'))
            for ext in extents:
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                width_cm = round(cx / 914400 * 2.54, 2)
                height_cm = round(cy / 914400 * 2.54, 2)
                images.append({
                    "width_cm": width_cm,
                    "height_cm": height_cm,
                })

        # Kiểm tra cả inline images (pict)
        pict_elems = run._element.findall(qn('w:pict'))
        if pict_elems:
            images.append({"type": "legacy_image"})

    return images


def extract_table_info(table):
    """Trích xuất thông tin chi tiết của bảng."""
    info = {
        "rows": len(table.rows),
        "cols": len(table.columns),
        "cells": [],
    }

    # Table alignment
    try:
        tbl_pr = table._tbl.find(qn('w:tblPr'))
        if tbl_pr is not None:
            jc = tbl_pr.find(qn('w:jc'))
            if jc is not None:
                info["alignment"] = jc.get(qn('w:val'), 'left')
    except Exception:
        pass

    # Table style
    if table.style and table.style.name:
        info["style"] = table.style.name

    # Analyze cells
    for row_idx, row in enumerate(table.rows):
        row_data = []
        for col_idx, cell in enumerate(row.cells):
            cell_info = {
                "row": row_idx,
                "col": col_idx,
                "text": cell.text.strip()[:100],  # Giới hạn 100 ký tự
            }

            # Cell shading (background color)
            tc = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is not None:
                shd = tcPr.find(qn('w:shd'))
                if shd is not None:
                    fill = shd.get(qn('w:fill'))
                    if fill and fill != 'auto':
                        cell_info["background"] = f"#{fill}"

                # Cell width
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    w = tcW.get(qn('w:w'))
                    w_type = tcW.get(qn('w:type'))
                    if w:
                        if w_type == 'dxa':
                            cell_info["width"] = f"{twips_to_cm(int(w))} cm"
                        elif w_type == 'pct':
                            cell_info["width"] = f"{int(w)/50}%"

                # Vertical merge
                vmerge = tcPr.find(qn('w:vMerge'))
                if vmerge is not None:
                    val = vmerge.get(qn('w:val'), 'continue')
                    cell_info["vmerge"] = val

                # Cell vertical alignment
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is not None:
                    cell_info["valign"] = vAlign.get(qn('w:val'))

            # Cell paragraph formatting
            if cell.paragraphs:
                first_para = cell.paragraphs[0]
                if first_para.alignment:
                    cell_info["text_align"] = get_alignment_name(first_para.alignment)
                if first_para.runs:
                    run_fmt = extract_run_formatting(first_para.runs[0])
                    if run_fmt:
                        cell_info["font_info"] = format_run_info(run_fmt)

            row_data.append(cell_info)
        info["cells"].append(row_data)

    return info


def extract_section_info(section):
    """Trích xuất thông tin page setup từ section."""
    info = {}

    # Page size
    if section.page_width:
        info["page_width"] = f"{emu_to_cm(section.page_width)} cm"
    if section.page_height:
        info["page_height"] = f"{emu_to_cm(section.page_height)} cm"

    # Orientation
    try:
        from docx.enum.section import WD_ORIENT
        if section.orientation == WD_ORIENT.LANDSCAPE:
            info["orientation"] = "Ngang (Landscape)"
        else:
            info["orientation"] = "Dọc (Portrait)"
    except Exception:
        info["orientation"] = "Dọc (Portrait)"

    # Margins
    if section.top_margin:
        info["margin_top"] = f"{emu_to_cm(section.top_margin)} cm"
    if section.bottom_margin:
        info["margin_bottom"] = f"{emu_to_cm(section.bottom_margin)} cm"
    if section.left_margin:
        info["margin_left"] = f"{emu_to_cm(section.left_margin)} cm"
    if section.right_margin:
        info["margin_right"] = f"{emu_to_cm(section.right_margin)} cm"

    # Gutter
    if section.gutter:
        info["gutter"] = f"{emu_to_cm(section.gutter)} cm"

    return info


def extract_styles_used(doc):
    """Trích xuất danh sách tất cả styles được sử dụng trong tài liệu."""
    styles_dict = {}

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        if style_name not in styles_dict:
            style_info = {"name": style_name, "count": 0, "type": "paragraph"}

            # Lấy thông tin style từ document styles
            try:
                style = doc.styles[style_name]
                if style.font.name:
                    style_info["font"] = style.font.name
                if style.font.size:
                    style_info["size"] = f"{style.font.size.pt}pt"
                if style.font.bold:
                    style_info["bold"] = True
                if style.font.italic:
                    style_info["italic"] = True
                if style.font.underline:
                    style_info["underline"] = True
                color = parse_color(style.font.color)
                if color:
                    style_info["color"] = color

                # Paragraph formatting from style
                pf = style.paragraph_format
                if pf.alignment:
                    style_info["alignment"] = get_alignment_name(pf.alignment)
                if pf.space_before:
                    style_info["space_before"] = f"{pf.space_before.pt}pt"
                if pf.space_after:
                    style_info["space_after"] = f"{pf.space_after.pt}pt"
                if pf.line_spacing:
                    style_info["line_spacing"] = str(pf.line_spacing)
                if pf.first_line_indent:
                    style_info["first_line_indent"] = f"{emu_to_cm(pf.first_line_indent)}cm"

            except (KeyError, AttributeError):
                pass

            # Lấy thông tin từ run đầu tiên (overrides)
            if para.runs:
                run_fmt = extract_run_formatting(para.runs[0])
                if run_fmt.get("font") and "font" not in style_info:
                    style_info["font"] = run_fmt["font"]
                if run_fmt.get("size") and "size" not in style_info:
                    style_info["size"] = run_fmt["size"]

            styles_dict[style_name] = style_info

        styles_dict[style_name]["count"] = styles_dict[style_name].get("count", 0) + 1

    return styles_dict


def extract_metadata(doc):
    """Trích xuất metadata của tài liệu."""
    props = doc.core_properties
    meta = {}

    if props.author:
        meta["Tác giả (Author)"] = props.author
    if props.title:
        meta["Tiêu đề (Title)"] = props.title
    if props.subject:
        meta["Chủ đề (Subject)"] = props.subject
    if props.keywords:
        meta["Từ khóa (Keywords)"] = props.keywords
    if props.category:
        meta["Phân loại (Category)"] = props.category
    if props.comments:
        meta["Ghi chú (Comments)"] = props.comments
    if props.last_modified_by:
        meta["Người sửa cuối (Last Modified By)"] = props.last_modified_by
    if props.created:
        meta["Ngày tạo (Created)"] = props.created.strftime("%Y-%m-%d %H:%M:%S")
    if props.modified:
        meta["Ngày sửa (Modified)"] = props.modified.strftime("%Y-%m-%d %H:%M:%S")
    if props.revision:
        meta["Phiên bản (Revision)"] = str(props.revision)

    return meta


def generate_markdown_template(doc, filename="document.docx"):
    """
    Hàm chính: Phân tích toàn bộ file .docx và tạo Markdown template chi tiết.
    """
    lines = []

    # ═══════════════════════════════════════════════════
    # TIÊU ĐỀ
    # ═══════════════════════════════════════════════════
    lines.append(f"# 📄 TEMPLATE PHÂN TÍCH TÀI LIỆU: `{filename}`")
    lines.append("")
    lines.append(f"> File được phân tích lúc: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("> Mục đích: Dùng làm RULE/GUIDELINE để sinh lại file .docx chuẩn format.")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ═══════════════════════════════════════════════════
    # PHẦN 1: METADATA
    # ═══════════════════════════════════════════════════
    lines.append("# PHẦN 1: THÔNG TIN TÀI LIỆU (METADATA)")
    lines.append("")

    metadata = extract_metadata(doc)
    if metadata:
        lines.append("| Thuộc tính | Giá trị |")
        lines.append("|---|---|")
        for key, val in metadata.items():
            lines.append(f"| {key} | {val} |")
    else:
        lines.append("*Không có metadata.*")
    lines.append("")

    # File name
    lines.append(f"- **Tên file gốc**: `{filename}`")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ═══════════════════════════════════════════════════
    # PHẦN 2: PAGE SETUP
    # ═══════════════════════════════════════════════════
    lines.append("# PHẦN 2: CÀI ĐẶT TRANG (PAGE SETUP)")
    lines.append("")

    for idx, section in enumerate(doc.sections):
        sec_info = extract_section_info(section)
        if len(doc.sections) > 1:
            lines.append(f"## Section {idx + 1}")
            lines.append("")

        lines.append("| Thuộc tính | Giá trị |")
        lines.append("|---|---|")
        for key, val in sec_info.items():
            label = {
                "page_width": "Chiều rộng trang",
                "page_height": "Chiều cao trang",
                "orientation": "Hướng giấy",
                "margin_top": "Lề trên (Top Margin)",
                "margin_bottom": "Lề dưới (Bottom Margin)",
                "margin_left": "Lề trái (Left Margin)",
                "margin_right": "Lề phải (Right Margin)",
                "gutter": "Gutter",
            }.get(key, key)
            lines.append(f"| {label} | {val} |")
        lines.append("")

    lines.append("---")
    lines.append("")

    # ═══════════════════════════════════════════════════
    # PHẦN 3: STYLES SỬ DỤNG
    # ═══════════════════════════════════════════════════
    lines.append("# PHẦN 3: QUY ĐỊNH STYLES (ĐỊNH DẠNG)")
    lines.append("")
    lines.append("> Danh sách tất cả các styles được sử dụng trong tài liệu kèm thuộc tính định dạng chi tiết.")
    lines.append("")

    styles = extract_styles_used(doc)

    # Sort: headings first, then by count
    def style_sort_key(item):
        name = item[0].lower()
        if 'heading' in name:
            # Extract heading number
            nums = re.findall(r'\d+', name)
            return (0, int(nums[0]) if nums else 99)
        if name == 'normal':
            return (1, 0)
        if name == 'title':
            return (0, -1)
        return (2, -item[1].get('count', 0))

    sorted_styles = sorted(styles.items(), key=style_sort_key)

    for style_name, style_info in sorted_styles:
        count = style_info.get("count", 0)
        lines.append(f"### Style: `{style_name}` (sử dụng {count} lần)")
        lines.append("")
        lines.append("| Thuộc tính | Giá trị |")
        lines.append("|---|---|")

        props_map = {
            "font": "Font chữ",
            "size": "Cỡ chữ",
            "bold": "In đậm",
            "italic": "In nghiêng",
            "underline": "Gạch chân",
            "color": "Màu chữ",
            "alignment": "Căn lề",
            "space_before": "Khoảng cách trước",
            "space_after": "Khoảng cách sau",
            "line_spacing": "Giãn dòng",
            "first_line_indent": "Thụt đầu dòng",
        }

        for prop_key, prop_label in props_map.items():
            if prop_key in style_info:
                val = style_info[prop_key]
                if isinstance(val, bool):
                    val = "Có" if val else "Không"
                lines.append(f"| {prop_label} | {val} |")

        lines.append("")

    lines.append("---")
    lines.append("")

    # ═══════════════════════════════════════════════════
    # PHẦN 4: CẤU TRÚC NỘI DUNG CHI TIẾT
    # ═══════════════════════════════════════════════════
    lines.append("# PHẦN 4: CẤU TRÚC NỘI DUNG CHI TIẾT")
    lines.append("")
    lines.append("> Mô tả tuần tự từng phần tử trong tài liệu: heading, paragraph, table, image, list...")
    lines.append("")

    element_index = 0
    table_index = 0
    current_list_id = None

    # Iterate through document body elements
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        # ── PARAGRAPH ──
        if tag == 'p':
            element_index += 1
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break

            if para is None:
                continue

            text = para.text.strip()
            style_name = para.style.name if para.style else "Normal"

            # Check if it's a list item
            list_info = detect_list_info(para)

            # Check for images
            images = detect_images_in_paragraph(para)

            # ── Heading ──
            if style_name.startswith("Heading") or style_name == "Title":
                level_match = re.search(r'\d+', style_name)
                level = int(level_match.group()) if level_match else 0
                md_heading = "#" * (level + 1) if level > 0 else "#"

                lines.append(f"---")
                lines.append("")
                lines.append(f"{md_heading} [{element_index}] {text}")
                lines.append("")
                lines.append(f"- **Style**: `{style_name}`")

                # Paragraph formatting
                para_fmt = extract_paragraph_formatting(para)
                for k, v in para_fmt.items():
                    label = {
                        "alignment": "Căn lề",
                        "left_indent": "Thụt trái",
                        "right_indent": "Thụt phải",
                        "first_line_indent": "Thụt đầu dòng",
                        "space_before": "Khoảng cách trước",
                        "space_after": "Khoảng cách sau",
                        "line_spacing": "Giãn dòng",
                    }.get(k, k)
                    if isinstance(v, bool):
                        v = "Có"
                    lines.append(f"- **{label}**: {v}")

                # Run formatting
                if para.runs:
                    unique_fmts = []
                    for run in para.runs:
                        fmt = extract_run_formatting(run)
                        fmt_str = format_run_info(fmt)
                        if fmt_str not in unique_fmts:
                            unique_fmts.append(fmt_str)
                    if unique_fmts:
                        lines.append(f"- **Định dạng text**: {' | '.join(unique_fmts)}")
                lines.append("")

            # ── List item ──
            elif list_info:
                indent = "  " * list_info["level"]
                marker = f"{list_info['type']}" 
                if list_info["type"] == "bullet":
                    lines.append(f"{indent}- [{element_index}] `[Bullet List, Level {list_info['level']}]` {text}")
                else:
                    lines.append(f"{indent}1. [{element_index}] `[Numbered List, Level {list_info['level']}]` {text}")

                # Show run formatting nếu khác default
                if para.runs:
                    fmt = extract_run_formatting(para.runs[0])
                    if fmt:
                        lines.append(f"{indent}   - Định dạng: {format_run_info(fmt)}")

            # ── Normal paragraph ──
            elif text:
                lines.append(f"**[{element_index}]** `Style: {style_name}` — {text[:200]}{'...' if len(text) > 200 else ''}")
                lines.append("")

                # Paragraph formatting
                para_fmt = extract_paragraph_formatting(para)
                fmt_parts = []
                for k, v in para_fmt.items():
                    label = {
                        "alignment": "Căn lề",
                        "left_indent": "Thụt trái",
                        "right_indent": "Thụt phải",
                        "first_line_indent": "Thụt đầu dòng",
                        "space_before": "Khoảng cách trước",
                        "space_after": "Khoảng cách sau",
                        "line_spacing": "Giãn dòng",
                    }.get(k, k)
                    if isinstance(v, bool):
                        v = "Có"
                    fmt_parts.append(f"{label}: {v}")
                if fmt_parts:
                    lines.append(f"  - Paragraph format: {' | '.join(fmt_parts)}")

                # Inline formatting changes
                if len(para.runs) > 1:
                    run_details = []
                    for r_idx, run in enumerate(para.runs):
                        if run.text.strip():
                            fmt = extract_run_formatting(run)
                            if fmt:
                                run_details.append(f'    - Run {r_idx+1}: "{run.text.strip()[:50]}" → {format_run_info(fmt)}')
                    if run_details:
                        lines.append(f"  - **Chi tiết Runs (inline formatting):**")
                        lines.extend(run_details)

                lines.append("")

            # ── Images in paragraph ──
            if images:
                for img in images:
                    if "width_cm" in img:
                        lines.append(f"  📷 `[Hình ảnh: {img['width_cm']}cm x {img['height_cm']}cm]`")
                    else:
                        lines.append(f"  📷 `[Hình ảnh (legacy format)]`")
                lines.append("")

        # ── TABLE ──
        elif tag == 'tbl':
            element_index += 1
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_info = extract_table_info(table)
                table_index += 1

                lines.append(f"---")
                lines.append("")
                lines.append(f"### [{element_index}] 📊 BẢNG (Table)")
                lines.append("")
                lines.append(f"- **Số hàng**: {table_info['rows']}")
                lines.append(f"- **Số cột**: {table_info['cols']}")
                if "style" in table_info:
                    lines.append(f"- **Table Style**: `{table_info['style']}`")
                if "alignment" in table_info:
                    lines.append(f"- **Căn lề bảng**: {table_info['alignment']}")
                lines.append("")

                # Render table structure
                if table_info["cells"]:
                    # Header row
                    header_row = table_info["cells"][0]
                    header_texts = [c.get("text", "")[:30] for c in header_row]

                    lines.append("**Cấu trúc bảng:**")
                    lines.append("")
                    lines.append("| " + " | ".join([f"Cột {i+1}" for i in range(len(header_texts))]) + " |")
                    lines.append("| " + " | ".join(["---"] * len(header_texts)) + " |")

                    for row_idx, row in enumerate(table_info["cells"]):
                        row_texts = []
                        for cell in row:
                            cell_text = cell.get("text", "")[:30]
                            if not cell_text:
                                cell_text = "(trống)"
                            row_texts.append(cell_text)
                        lines.append("| " + " | ".join(row_texts) + " |")

                    lines.append("")

                    # Cell formatting details
                    lines.append("**Chi tiết định dạng cells:**")
                    lines.append("")
                    for row_idx, row in enumerate(table_info["cells"][:3]):  # Chỉ show 3 rows đầu
                        for cell in row:
                            details = []
                            if "background" in cell:
                                details.append(f"Background: {cell['background']}")
                            if "width" in cell:
                                details.append(f"Width: {cell['width']}")
                            if "valign" in cell:
                                details.append(f"VAlign: {cell['valign']}")
                            if "vmerge" in cell:
                                details.append(f"VMerge: {cell['vmerge']}")
                            if "font_info" in cell:
                                details.append(f"Font: {cell['font_info']}")
                            if "text_align" in cell:
                                details.append(f"Align: {cell['text_align']}")
                            if details:
                                lines.append(f"  - Cell [{cell['row']},{cell['col']}]: {' | '.join(details)}")

                    if len(table_info["cells"]) > 3:
                        lines.append(f"  - *(... còn {len(table_info['cells']) - 3} hàng nữa)*")

                lines.append("")

    lines.append("---")
    lines.append("")

    # ═══════════════════════════════════════════════════
    # PHẦN 5: TÓM TẮT & HƯỚNG DẪN SỬ DỤNG
    # ═══════════════════════════════════════════════════
    lines.append("# PHẦN 5: TÓM TẮT & HƯỚNG DẪN SỬ DỤNG")
    lines.append("")
    lines.append("## Thống kê")
    lines.append("")

    total_paragraphs = len(doc.paragraphs)
    total_tables = len(doc.tables)
    total_headings = sum(1 for p in doc.paragraphs if p.style and p.style.name.startswith("Heading"))
    total_images = sum(len(detect_images_in_paragraph(p)) for p in doc.paragraphs)

    lines.append(f"- Tổng số paragraphs: **{total_paragraphs}**")
    lines.append(f"- Tổng số headings: **{total_headings}**")
    lines.append(f"- Tổng số bảng: **{total_tables}**")
    lines.append(f"- Tổng số hình ảnh: **{total_images}**")
    lines.append(f"- Tổng số styles sử dụng: **{len(styles)}**")
    lines.append("")

    lines.append("## Hướng dẫn sử dụng Template này")
    lines.append("")
    lines.append("1. **Tạo file .docx mới** với cùng Page Setup (khổ giấy, margins) như mô tả ở Phần 2.")
    lines.append("2. **Thiết lập Styles** theo đúng các thuộc tính ghi ở Phần 3 (font, size, color, spacing...).")
    lines.append("3. **Xây dựng nội dung** theo đúng cấu trúc tuần tự ở Phần 4.")
    lines.append("4. **Lưu ý**: Các Heading levels, bullet/numbered lists, và bảng phải giữ đúng format.")
    lines.append("5. **Hình ảnh**: Chèn lại hình ảnh đúng kích thước và vị trí như mô tả.")
    lines.append("")
    lines.append("---")
    lines.append("*File template này được tạo tự động bởi Doc2MD Template Analyzer.*")

    return "\n".join(lines)


def parse_docx_to_md(file_path: str) -> str:
    """
    Entry point: Đọc file .docx và trả về nội dung Markdown.
    """
    doc = Document(file_path)
    filename = os.path.basename(file_path)
    return generate_markdown_template(doc, filename)
