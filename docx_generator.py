"""
docx_generator.py
Module tạo file .docx từ nội dung Markdown SRS.
Được tích hợp vào API endpoint /api/generate-docx.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colors ──
HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BG = "D9D9D9"


# ════════════════════════════════════
# XML HELPERS
# ════════════════════════════════════

def _set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)


def _cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcPr.append(tcMar)


# ════════════════════════════════════
# PARAGRAPH HELPERS
# ════════════════════════════════════

def _add_run(para, text, bold=False, italic=False, size=11, color=BLACK, underline=False):
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.underline = underline
    return run


def _add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
              italic=False, size=11, before=0, after=4):
    para = doc.add_paragraph()
    if text:
        _add_run(para, text, bold=bold, italic=italic, size=size)
    para.alignment = align
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)
    return para


def _apply_heading(doc, text, level):
    style_name = f'Heading {level}'
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    sizes = {1: 14, 2: 12, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(18 if level == 1 else 12 if level == 2 else 6)
    pf.space_after = Pt(6)
    return para


def _add_list(doc, text, level=0):
    para = doc.add_paragraph(style='List Paragraph')
    _add_run(para, text)
    pf = para.paragraph_format
    pf.left_indent = Cm(0.63 * (level + 1))
    pf.first_line_indent = Cm(-0.63)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return para


# ════════════════════════════════════
# TABLE BUILDER
# ════════════════════════════════════

def _make_table(doc, headers, rows, col_widths=None, header_bg=HEADER_BG):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        _add_run(cell.paragraphs[0], h, bold=True)
        set_cell_shading = _set_cell_shading
        set_cell_shading(cell, header_bg)
        _cell_margins(cell)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            _add_run(cell.paragraphs[0], str(val) if val else '')
            _cell_margins(cell)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    _set_table_borders(table)
    doc.add_paragraph()
    return table


# ════════════════════════════════════
# MARKDOWN PARSER → DOCX ELEMENTS
# ════════════════════════════════════

def _parse_markdown_to_docx(doc, md_text):
    """
    Parse markdown content và thêm vào document.
    Xử lý: headings, tables, lists, bold/italic inline, normal paragraphs.
    """
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Heading detection: # ## ###
        h_match = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2).strip()
            _apply_heading(doc, text, level)
            i += 1
            continue

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            i += 1
            continue

        # Table detection: starts with |
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # Skip separator rows like |---|---|
                if not re.match(r'^\|[\s\-:|]+\|$', row_line):
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    table_lines.append(cells)
                i += 1
            if len(table_lines) >= 1:
                headers = table_lines[0]
                rows = table_lines[1:]
                _make_table(doc, headers, rows)
            continue

        # List items: - or numbered
        ul_match = re.match(r'^(\s*)-\s+(.+)$', line)
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ul_match:
            level = len(ul_match.group(1)) // 2
            _add_list(doc, _strip_inline(ul_match.group(2)), level=level)
            i += 1
            continue
        if ol_match:
            level = len(ol_match.group(1)) // 2
            _add_list(doc, _strip_inline(ol_match.group(2)), level=level)
            i += 1
            continue

        # Blockquote: > text
        if stripped.startswith('> '):
            text = stripped[2:]
            p = _add_para(doc, _strip_inline(text), italic=True, after=4)
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        _render_inline(para, stripped)
        para.paragraph_format.space_after = Pt(4)
        i += 1


def _strip_inline(text):
    """Remove markdown bold/italic markers for simple text."""
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text


def _render_inline(para, text):
    """Render inline markdown (bold, italic) as Word runs."""
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|(.+?)(?=\*\*\*|\*\*|\*|`|$))', re.DOTALL)
    pos = 0
    while pos < len(text):
        # bold+italic
        m = re.match(r'\*\*\*(.+?)\*\*\*', text[pos:])
        if m:
            _add_run(para, m.group(1), bold=True, italic=True)
            pos += m.end()
            continue
        # bold
        m = re.match(r'\*\*(.+?)\*\*', text[pos:])
        if m:
            _add_run(para, m.group(1), bold=True)
            pos += m.end()
            continue
        # italic
        m = re.match(r'\*(.+?)\*', text[pos:])
        if m:
            _add_run(para, m.group(1), italic=True)
            pos += m.end()
            continue
        # code
        m = re.match(r'`(.+?)`', text[pos:])
        if m:
            _add_run(para, m.group(1))
            pos += m.end()
            continue
        # plain text until next marker
        m = re.match(r'.+?(?=\*\*\*|\*\*|\*|`|$)', text[pos:], re.DOTALL)
        if m:
            _add_run(para, m.group(0))
            pos += m.end()
        else:
            _add_run(para, text[pos:])
            break


# ════════════════════════════════════
# PUBLIC API: build_srs_docx
# ════════════════════════════════════

def build_srs_docx(title: str = "SRS Document", markdown_content: str = "") -> bytes:
    """
    Build a .docx file from markdown content.
    Returns the file as bytes for streaming response.
    
    Args:
        title: document title (used for cover page)
        markdown_content: full markdown text to convert
    
    Returns:
        bytes: .docx file content
    """
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # ── Default style ──
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK

    # ── Override Heading styles ──
    for lvl, sz in [(1, 14), (2, 12), (3, 11)]:
        try:
            h = doc.styles[f'Heading {lvl}']
            h.font.name = 'Calibri'
            h.font.size = Pt(sz)
            h.font.color.rgb = HEADING_COLOR
            h.font.bold = True
            h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    # ── Cover Page ──
    for _ in range(3):
        doc.add_paragraph()
    _add_para(doc, title, WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, after=6)
    _add_para(doc, "Đặc tả yêu cầu hệ thống (System Requirement Specification)",
              WD_ALIGN_PARAGRAPH.CENTER, size=14, after=6)
    _add_para(doc, "Phiên bản: 1.0", WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    _add_para(doc, "Hà Nội, Tháng 03/2026", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=4)
    doc.add_page_break()

    # ── Parse and render markdown content ──
    if markdown_content.strip():
        _parse_markdown_to_docx(doc, markdown_content)

    # ── Save to bytes ──
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
