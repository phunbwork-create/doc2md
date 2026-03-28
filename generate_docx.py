"""
Generate SRS .docx — Fixed version.
Fixes: (1) Use Word built-in Heading styles, (2) Add table borders, (3) Fix alignment
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color constants ──
HEADING_COLOR = RGBColor(0x1F, 0x38, 0x64)   # #1F3864 dark navy
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BG = "D9D9D9"
LIGHT_BG = "E7E6E6"


# ═══════════════════════════════════════════════════
# XML HELPERS
# ═══════════════════════════════════════════════════

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shd
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)


def set_table_borders(table):
    """Add visible borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove old tblBorders
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    borders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{border_name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)


def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    # Remove existing
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcPr.append(tcMar)


# ═══════════════════════════════════════════════════
# HEADING using Word built-in styles
# ═══════════════════════════════════════════════════

def apply_heading_style(doc, text, level):
    """Add a heading using Word's built-in Heading 1/2/3 styles, then override font."""
    style_name = f'Heading {level}'
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.color.rgb = HEADING_COLOR
    run.bold = True
    sizes = {1: 14, 2: 12, 3: 11}
    run.font.size = Pt(sizes.get(level, 11))
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        pf.space_before = Pt(18)
        pf.space_after = Pt(6)
    elif level == 2:
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
    elif level == 3:
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
    return para


# ═══════════════════════════════════════════════════
# PARAGRAPH HELPERS
# ═══════════════════════════════════════════════════

def add_para(doc, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False,
             italic=False, size=11, color=BLACK, underline=False,
             before=0, after=4):
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = color
        run.underline = underline
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    return para


def add_run(para, text, bold=False, italic=False, size=11,
            color=BLACK, underline=False):
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.underline = underline
    return run


def add_list_para(doc, text, level=0, size=11):
    para = doc.add_paragraph(style='List Paragraph')
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    pf = para.paragraph_format
    pf.left_indent = Cm(0.63 * (level + 1))
    pf.first_line_indent = Cm(-0.63)
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    return para


# ═══════════════════════════════════════════════════
# TABLE BUILDER
# ═══════════════════════════════════════════════════

def make_table(doc, headers, rows, col_widths_cm=None, header_bg=HEADER_BG):
    ncols = len(headers)
    nrows = 1 + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'  # gives default borders

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = BLACK
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, header_bg)
        set_cell_margins(cell)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        tr = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = tr.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val) if val else '')
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_margins(cell)

    # Column widths
    if col_widths_cm:
        for row in table.rows:
            for i, w in enumerate(col_widths_cm):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)

    set_table_borders(table)
    doc.add_paragraph()  # spacing after table
    return table


def make_uc_table(doc, rows_data):
    """2-column UC info table: label (grey) | value"""
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(rows_data):
        c0, c1 = table.rows[i].cells[0], table.rows[i].cells[1]
        c0.text = ''
        r0 = c0.paragraphs[0].add_run(label)
        r0.font.name = 'Calibri'; r0.font.size = Pt(11)
        r0.bold = True; r0.font.color.rgb = BLACK
        set_cell_shading(c0, HEADER_BG)
        set_cell_margins(c0)
        c0.width = Cm(4.03)

        c1.text = ''
        r1 = c1.paragraphs[0].add_run(value)
        r1.font.name = 'Calibri'; r1.font.size = Pt(11)
        r1.font.color.rgb = BLACK
        set_cell_margins(c1)
        c1.width = Cm(13.77)

    set_table_borders(table)
    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════
# MAIN BUILD
# ═══════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # ── Default Normal style ──
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK

    # ── Override Heading styles globally ──
    for lvl, sz in [(1, 14), (2, 12), (3, 11)]:
        try:
            h_style = doc.styles[f'Heading {lvl}']
            h_style.font.name = 'Calibri'
            h_style.font.size = Pt(sz)
            h_style.font.color.rgb = HEADING_COLOR
            h_style.font.bold = True
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass

    # ════════════════════════════════════════
    # TRANG BÌA
    # ════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    t = add_para(doc, "ABC CORPORATION", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, before=6, after=6)
    add_para(doc, "HỆ THỐNG QUẢN LÝ KHÁCH HÀNG (CRM)", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, after=6)
    add_para(doc, "Đặc tả yêu cầu hệ thống", WD_ALIGN_PARAGRAPH.CENTER, size=14, after=4)
    add_para(doc, "(System Requirement Specification)", WD_ALIGN_PARAGRAPH.CENTER, size=14, after=12)
    add_para(doc, "Phiên bản: 1.0", WD_ALIGN_PARAGRAPH.CENTER, size=12, after=4)
    add_para(doc, "Hà Nội, Tháng 03/2026", WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12, after=4)
    doc.add_page_break()

    # ════════════════════════════════════════
    # APPROVAL PAGE (plain table, no heading style)
    # ════════════════════════════════════════
    add_para(doc, "Approval Page", bold=True, size=13, after=8)
    add_para(doc,
        'Việc xác nhận trên tài liệu này bởi đại diện ủy quyền của ABC Corporation cho thấy sự đồng ý giữa ABC Corporation và đơn vị phát triển về tài liệu "CRM – Đặc tả yêu cầu hệ thống".',
        align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=12)

    apr_table = doc.add_table(rows=8, cols=4)
    apr_table.style = 'Table Grid'
    apr_data = [
        ("Người soạn:", "(Đơn vị PT)\nBusiness Analyst", "Chữ ký:", "____________________"),
        ("",            "",                                "Ngày:",   "____/____/____"),
        ("Người review:", "(Đơn vị PT)\nProject Manager", "Chữ ký:", "____________________"),
        ("",            "",                                "Ngày:",   "____/____/____"),
        ("Hỗ trợ bởi:", "(ABC Corporation)",              "Chữ ký:", "____________________"),
        ("",            "",                                "Ngày:",   "____/____/____"),
        ("Phê duyệt bởi:", "(ABC Corporation)",           "Chữ ký:", "____________________"),
        ("",            "",                                "Ngày:",   "____/____/____"),
    ]
    col_w = [3.0, 5.0, 2.5, 5.0]
    for r_idx, row_d in enumerate(apr_data):
        for c_idx, val in enumerate(row_d):
            cell = apr_table.rows[r_idx].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            is_bold = c_idx == 0 and val not in ("", )
            run = p.add_run(val)
            run.font.name = 'Calibri'; run.font.size = Pt(11)
            run.bold = is_bold
            run.font.color.rgb = BLACK
            set_cell_margins(cell)
            cell.width = Cm(col_w[c_idx])
    set_table_borders(apr_table)
    doc.add_paragraph()
    doc.add_page_break()

    # ════════════════════════════════════════
    # REVISION HISTORY
    # ════════════════════════════════════════
    add_para(doc, "Revision History", bold=True, size=13, after=8)
    make_table(doc,
        ["Ngày", "Phiên bản", "Tác giả", "Mô tả thay đổi"],
        [
            ["26/03/2026", "0.5",   "Nguyễn Văn A", "Tạo mới lần đầu"],
            ["",           "0.7",   "",              "Cho review nội bộ"],
            ["",           "0.8",   "",              "Phát hành lần đầu tới khách hàng"],
            ["",           "0.9.x", "",              "Cập nhật theo review khách hàng"],
            ["",           "1.0",   "",              "Cho ký duyệt chính thức"],
        ],
        col_widths_cm=[2.75, 2.07, 3.3, 8.38]
    )
    doc.add_page_break()

    # ════════════════════════════════════════
    # 1. GIỚI THIỆU
    # ════════════════════════════════════════
    apply_heading_style(doc, "1. Giới thiệu", 1)

    apply_heading_style(doc, "1.1. Mục đích", 2)
    add_para(doc, "Tài liệu Đặc tả yêu cầu chức năng sẽ:", after=4)
    for item in [
        "Xác định phạm vi mục tiêu kinh doanh, chức năng nghiệp vụ và các đơn vị tổ chức được bao phủ.",
        "Xác định các quy trình nghiệp vụ mà giải pháp cần hỗ trợ.",
        "Tạo sự hiểu biết chung về các yêu cầu chức năng cho tất cả các bên liên quan.",
        "Thiết lập cơ sở để xác định các bài kiểm tra chấp nhận (Acceptance Test) cho giải pháp.",
    ]:
        add_list_para(doc, item)
    add_para(doc, "", after=4)
    add_para(doc,
        "Mục đích của tài liệu này là thu thập và phân tích tất cả các ý tưởng để xác định hệ thống, các yêu cầu liên quan đến người dùng.",
        after=4)

    apply_heading_style(doc, "1.2. Tổng quan", 2)
    add_para(doc,
        "Hệ thống Quản lý Khách hàng (CRM) là một nền tảng web cho phép doanh nghiệp quản lý toàn diện thông tin khách hàng, theo dõi lịch sử tương tác, quản lý cơ hội kinh doanh và tối ưu hóa quy trình chăm sóc khách hàng. Hệ thống hỗ trợ nhiều vai trò người dùng: Nhân viên kinh doanh, Quản lý, và Quản trị hệ thống.",
        after=4)

    apply_heading_style(doc, "1.3. Đối tượng đọc và gợi ý", 2)
    add_para(doc, "Tài liệu này dành cho:", after=4)
    for item in [
        "Đội phát triển: Chịu trách nhiệm thiết kế chi tiết, lập trình và thực hiện kiểm thử.",
        "Đội di chuyển dữ liệu: Chịu trách nhiệm tạo script di chuyển dữ liệu.",
        "Đội tài liệu: Chịu trách nhiệm viết Hướng dẫn sử dụng cho ứng dụng.",
        "Đội UAT: Chịu trách nhiệm tiến hành các phiên kiểm thử chấp nhận người dùng.",
    ]:
        add_list_para(doc, item)
    add_para(doc, "", after=4)

    apply_heading_style(doc, "1.4. Thuật ngữ viết tắt", 2)
    make_table(doc,
        ["Viết tắt", "Tham chiếu"],
        [
            ["SRS",  "Đặc tả yêu cầu hệ thống (System Requirement Specification)"],
            ["UC",   "Trường hợp sử dụng (Use Case)"],
            ["BR",   "Quy tắc nghiệp vụ (Business Rules)"],
            ["CBR",  "Quy tắc nghiệp vụ chung (Common Business Rules)"],
            ["ET",   "Mẫu email (Email Template)"],
            ["N/A",  "Không áp dụng hoặc Không có sẵn"],
            ["MSG",  "Thông báo (Message)"],
            ["CRM",  "Quản lý quan hệ khách hàng (Customer Relationship Management)"],
            ["CSKH", "Chăm sóc khách hàng"],
        ],
        col_widths_cm=[3.04, 14.76]
    )

    apply_heading_style(doc, "1.5. Tài liệu tham chiếu", 2)
    make_table(doc,
        ["Tiêu đề", "Tham chiếu", "Mô tả"],
        [
            ["Tài liệu BRD – CRM", "Link nội bộ", "Tài liệu yêu cầu nghiệp vụ cho hệ thống CRM"],
            ["Tài liệu thiết kế UI/UX", "Figma Link", "Thiết kế giao diện người dùng cho CRM"],
        ],
        col_widths_cm=[4.26, 5.08, 8.46]
    )

    # ════════════════════════════════════════
    # 2. YÊU CẦU TỔNG QUAN
    # ════════════════════════════════════════
    apply_heading_style(doc, "2. Yêu cầu tổng quan", 1)
    add_para(doc, "Phần này mô tả tổng quan chức năng hệ thống hoặc các quy trình nghiệp vụ được thể hiện trong các sơ đồ khác nhau.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4)

    apply_heading_style(doc, "2.1. Sơ đồ ngữ cảnh (Context Diagram)", 2)
    add_para(doc, "Phần này cho thấy mối quan hệ tĩnh giữa các đối tượng trong hệ thống.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=4)
    add_para(doc, "[Hình ảnh: Sơ đồ ngữ cảnh hệ thống CRM]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    add_para(doc, "Hình 1: Sơ đồ ngữ cảnh hệ thống CRM", WD_ALIGN_PARAGRAPH.CENTER, after=10)
    add_para(doc, "Object Description:", bold=True, after=4)
    make_table(doc,
        ["#", "Đối tượng", "Mô tả"],
        [
            ["",  "Dữ liệu",               ""],
            ["1", "Khách hàng",             "Lưu trữ thông tin KH: họ tên, email, SĐT, địa chỉ, loại KH"],
            ["2", "Cơ hội kinh doanh",      "Lưu trữ thông tin về các cơ hội bán hàng tiềm năng"],
            ["",  "Tác nhân",              ""],
            ["1", "Nhân viên kinh doanh",   "Tạo, xem, cập nhật thông tin khách hàng"],
            ["2", "Quản lý",               "Phê duyệt, xem báo cáo, quản lý phân quyền"],
            ["",  "Hệ thống bên ngoài",   ""],
            ["1", "Hệ thống Email",         "Gửi thông báo email tự động"],
            ["2", "Hệ thống ERP",           "Đồng bộ thông tin khách hàng"],
        ],
        col_widths_cm=[1.17, 4.02, 12.6]
    )

    apply_heading_style(doc, "2.2. Quy trình nghiệp vụ (Workflow)", 2)
    add_para(doc, "Phần này mô tả luồng công việc và các bước mà người dùng thực hiện để hoàn thành quy trình nghiệp vụ.", after=4)
    add_para(doc, "[Hình ảnh: Sơ đồ quy trình]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    add_para(doc, "Hình 2: Quy trình quản lý khách hàng", WD_ALIGN_PARAGRAPH.CENTER, after=10)

    apply_heading_style(doc, "2.3. Sơ đồ chuyển trạng thái (State Transition)", 2)
    add_para(doc, "Sơ đồ này thể hiện hành vi của hệ thống khi phản hồi hành động của người dùng bằng cách thay đổi trạng thái của đối tượng.", after=4)
    add_para(doc, "[Hình ảnh: Sơ đồ trạng thái — Khách hàng]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    add_para(doc, "Hình 3: Sơ đồ chuyển trạng thái — Khách hàng", WD_ALIGN_PARAGRAPH.CENTER, after=10)

    apply_heading_style(doc, "2.4. Ma trận phân quyền (Permission Matrix)", 2)
    add_para(doc, "Ma trận phân quyền ánh xạ chức năng và vai trò người dùng:", after=4)
    p = add_para(doc, "", after=2)
    add_run(p, "Ghi chú: ", bold=True)
    add_run(p, '"O" = có quyền. "O*" = có quyền với mục tự tạo. "X" = không có quyền.')
    doc.add_paragraph()
    make_table(doc,
        ["Chức năng", "Nhân viên KD", "Quản lý", "Admin"],
        [
            ["Quản lý Khách hàng", "", "", ""],
            ["Tạo thông tin khách hàng",      "O",  "O", "O"],
            ["Xem thông tin khách hàng",       "O",  "O", "O"],
            ["Cập nhật thông tin khách hàng",  "O*", "O", "O"],
            ["Xóa thông tin khách hàng",       "X",  "O", "O"],
            ["Tìm kiếm khách hàng",            "O",  "O", "O"],
            ["Xuất báo cáo khách hàng",        "X",  "O", "O"],
        ],
        col_widths_cm=[6.5, 3.0, 3.0, 3.0]
    )

    apply_heading_style(doc, "2.5. Bản đồ trang (Site Map)", 2)
    add_para(doc, "Bản đồ trang mô tả cách điều hướng qua hệ thống.", after=4)
    make_table(doc,
        ["Trang", "Mô tả", "Phân quyền"],
        [
            ["Trang chủ (Dashboard)", "Trang tổng quan sau đăng nhập", "Tất cả người dùng"],
            ["Danh sách khách hàng", "Danh sách tất cả KH", "NV KD, Quản lý, Admin"],
            ["Chi tiết khách hàng", "Xem/chỉnh sửa thông tin KH", "NV KD, Quản lý, Admin"],
            ["Tạo khách hàng mới", "Nhập thông tin KH mới", "NV KD, Quản lý, Admin"],
        ],
        col_widths_cm=[4.0, 8.0, 5.8]
    )

    # ════════════════════════════════════════
    # 3. ĐẶC TẢ USE CASE
    # ════════════════════════════════════════
    apply_heading_style(doc, "3. Đặc tả Use Case", 1)
    add_para(doc, "Phần này bao gồm các yêu cầu chức năng chi tiết của hệ thống, mô tả đầu vào, hành vi và đầu ra mong đợi.", after=4)

    apply_heading_style(doc, "3.1. Quản lý Khách hàng", 2)
    apply_heading_style(doc, "UC 1: Tạo thông tin khách hàng", 3)

    make_uc_table(doc, [
        ("Mục tiêu (Objective):",
         "Cho phép người dùng tạo mới thông tin khách hàng trong hệ thống CRM bao gồm: thông tin cá nhân, thông tin liên hệ, phân loại khách hàng, và ghi chú."),
        ("Tác nhân (Actor):",
         'User có role = "Nhân viên kinh doanh" hoặc "Quản lý" hoặc "Quản trị hệ thống"'),
        ("Kích hoạt (Trigger):",
         'User click vào nút "Tạo khách hàng mới" trên màn hình Danh sách khách hàng'),
        ("Điều kiện tiên quyết (Pre-condition):",
         "User đăng nhập thành công vào hệ thống CRM và có quyền tạo khách hàng"),
        ("Điều kiện sau (Post-condition):",
         "Thông tin khách hàng mới được lưu thành công vào hệ thống"),
    ])

    add_para(doc, "Activities Flow", bold=True, after=4)
    add_para(doc, "[Hình ảnh: Sơ đồ luồng hoạt động UC1]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=10)

    add_para(doc, "User Story:", bold=True, after=4)
    add_para(doc, "US: Tạo thông tin khách hàng", bold=True, after=4)
    add_para(doc, "Với vai trò là nhân viên kinh doanh", after=4)
    add_para(doc, 'Tôi mong muốn có chức năng "Tạo thông tin khách hàng"', after=4)
    add_para(doc, "Để ghi nhận đầy đủ thông tin khách hàng tiềm năng vào hệ thống, phục vụ cho việc chăm sóc và theo dõi về sau.", after=8)

    add_para(doc, "AC (Acceptance Criteria):", bold=True, after=4)

    add_list_para(doc, 'AC 1: Khi user click "Tạo khách hàng mới", hệ thống hiển thị màn hình với các trường thông tin.')
    add_para(doc, "[Hình ảnh: Screenshot màn hình Tạo thông tin khách hàng]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)

    add_list_para(doc, 'AC 2: Khi user nhập thông tin và ấn "Lưu", hệ thống validate theo các logic sau:')
    add_list_para(doc, "[Họ và tên] không được để trống → MSG 1", level=1)
    add_list_para(doc, "[Email] phải đúng định dạng hợp lệ → MSG 2", level=1)
    add_list_para(doc, "[Số điện thoại] chỉ chấp nhận số, 10-11 ký tự → MSG 3", level=1)
    add_list_para(doc, "[Email] phải là duy nhất trong hệ thống → MSG 6", level=1)
    add_list_para(doc, "[Loại khách hàng] phải được chọn → MSG 1", level=1)

    add_list_para(doc, 'AC 3: Khi pass qua tất cả validation, hệ thống xử lý:')
    add_list_para(doc, 'Tạo bản ghi mới với trạng thái = "Hoạt động"', level=1)
    add_list_para(doc, "Sinh mã KH tự động: KH-YYYYMMDD-XXX", level=1)
    add_list_para(doc, "Ghi nhận người tạo và thời gian tạo", level=1)
    add_list_para(doc, "Hiển thị MSG 9 và redirect về màn hình Chi tiết KH", level=1)

    add_list_para(doc, 'AC 4: Khi user ấn "Hủy", hiển thị xác nhận MSG 4. Đồng ý → quay lại Danh sách. Hủy → ở lại.')
    doc.add_paragraph()

    add_para(doc, "Business Rules", bold=True, after=4)
    make_table(doc,
        ["Step", "BR Code", "Mô tả"],
        [
            ["(1)", "BR 1",
             "Quy tắc hiển thị màn hình:\nHiển thị các trường: Họ và tên (*), Email (*), Số điện thoại, Ngày sinh, Giới tính (Nam/Nữ/Khác), Địa chỉ, Loại KH (*) [Cá nhân/Doanh nghiệp], Tên công ty, Mã số thuế, Nguồn KH, Ghi chú.\nMặc định: [Loại KH] = \"Cá nhân\". Khi = \"Doanh nghiệp\" → hiển thị [Tên công ty] và [MST] (bắt buộc)."],
            ["(2)", "BR 2",
             "Quy tắc kiểm tra dữ liệu:\n(a) [Họ và tên]: bắt buộc, tối đa 200 ký tự.\n(b) [Email]: bắt buộc, đúng format, duy nhất, tối đa 100 ký tự.\n(c) [SĐT]: tùy chọn, chỉ số, 10-11 ký tự.\n(d) [Ngày sinh]: phải nhỏ hơn ngày hiện tại.\n(e) [Địa chỉ]: tối đa 500 ký tự.\n(f) [MST]: bắt buộc khi Loại KH=Doanh nghiệp, format 10 hoặc 13 số.\nNếu lỗi: hiển thị In-line Error màu đỏ bên dưới trường."],
            ["(3)", "BR 3",
             "Quy tắc lưu dữ liệu:\n(a) Tạo bản ghi mới; (b) Sinh [Mã KH] = KH-YYYYMMDD-XXX;\n(c) [Trạng thái] = \"Hoạt động\"; (d) [Ngày tạo] = thời điểm hiện tại;\n(e) [Người tạo] = user hiện tại; (f) Hiển thị MSG 9;\n(g) Redirect về màn hình Chi tiết KH vừa tạo."],
            ["(3)", "BR 4",
             "Quy tắc thông báo email:\nSau khi tạo KH thành công, hệ thống gửi email cho Quản lý theo mẫu ET 1 (nếu được cấu hình)."],
        ],
        col_widths_cm=[1.46, 1.78, 14.55]
    )

    apply_heading_style(doc, "3.2. Quy tắc nghiệp vụ chung (Common Business Rules)", 2)
    add_para(doc, "Phần này mô tả các quy tắc nghiệp vụ được áp dụng chung cho nhiều use case.", after=4)
    make_table(doc,
        ["BR Code", "Mô tả"],
        [["CBR1", "Quy tắc định dạng ngày tháng: Tất cả trường ngày tháng hiển thị theo format DD/MM/YYYY. Khi nhập liệu sử dụng Date Picker."]],
        col_widths_cm=[2.51, 15.29]
    )

    # ════════════════════════════════════════
    # 4. MOCKUP SCREENS
    # ════════════════════════════════════════
    apply_heading_style(doc, "4. Màn hình giao diện (Mockups Screen)", 1)
    add_para(doc, "Phần này chứa các màn hình và thuộc tính tương ứng, liên kết với một hoặc nhiều use case được mô tả ở trên.", after=4)

    apply_heading_style(doc, "4.1. Màn hình Danh sách khách hàng", 2)
    add_para(doc, "Mục đích: Hiển thị danh sách tất cả khách hàng.\nTruy cập: Menu trái > Khách hàng > Danh sách.", after=4)
    add_para(doc, "[Hình ảnh: Screenshot màn hình Danh sách khách hàng]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    add_para(doc, "Màn hình 1: Danh sách khách hàng", WD_ALIGN_PARAGRAPH.CENTER, after=10)

    make_table(doc,
        ["#", "Thành phần", "Giá trị", "Mô tả"],
        [
            ["1", "Ô tìm kiếm",         "Free text", "Tìm kiếm theo tên, email, SĐT. Sắp xếp mặc định: Ngày tạo giảm dần"],
            ["2", "Nút \"Tạo KH mới\"", "Button",    "Mở màn hình Tạo thông tin KH mới — tham chiếu UC 1"],
            ["3", "Bảng danh sách",      "Table",     "Hiển thị: Mã KH, Họ tên, Email, SĐT, Loại KH, Trạng thái, Ngày tạo"],
            ["4", "Nút \"Xuất Excel\"",  "Button",    "Xuất danh sách ra Excel. Chỉ hiển thị cho Quản lý và Admin"],
        ],
        col_widths_cm=[1.6, 4.5, 3.5, 8.2]
    )

    apply_heading_style(doc, "4.2. Màn hình Tạo thông tin khách hàng", 2)
    add_para(doc, 'Mục đích: Cho phép nhập thông tin khách hàng mới.\nTruy cập: Từ nút "Tạo KH mới" trên Màn hình Danh sách.', after=4)
    add_para(doc, "[Hình ảnh: Screenshot Tạo thông tin khách hàng]", WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=4)
    add_para(doc, "Màn hình 2: Tạo thông tin khách hàng", WD_ALIGN_PARAGRAPH.CENTER, after=10)

    make_table(doc,
        ["#", "Thành phần", "Loại", "Sửa được", "Bắt buộc", "Mặc định", "Mô tả"],
        [
            ["1",  "Họ và tên",      "Single Line",  "Có", "Có (*)",   "",          "Tối đa 200 ký tự"],
            ["2",  "Email",          "Single Line",  "Có", "Có (*)",   "",          "Phải duy nhất. Tối đa 100 ký tự"],
            ["3",  "Số điện thoại",  "Single Line",  "Có", "Không",    "",          "Chỉ nhận số, 10-11 ký tự"],
            ["4",  "Ngày sinh",      "Date Picker",  "Có", "Không",    "",          "Phải < ngày hiện tại. Format DD/MM/YYYY"],
            ["5",  "Giới tính",      "Radio Button", "Có", "Không",    "",          "Giá trị: Nam / Nữ / Khác"],
            ["6",  "Địa chỉ",        "Multi Line",   "Có", "Không",    "",          "Tối đa 500 ký tự"],
            ["7",  "Loại KH",        "Dropdown",     "Có", "Có (*)",   "Cá nhân",  "Cá nhân / Doanh nghiệp"],
            ["8",  "Tên công ty",    "Single Line",  "Có", "Có(*) DN", "",          "Chỉ hiển thị khi Loại KH = Doanh nghiệp"],
            ["9",  "Mã số thuế",     "Single Line",  "Có", "Có(*) DN", "",          "10 hoặc 13 số. Chỉ khi Doanh nghiệp"],
            ["10", "Nguồn KH",       "Dropdown",     "Có", "Không",    "",          "Website / Giới thiệu / QC / Sự kiện / Khác"],
            ["11", "Ghi chú",        "Multi Line",   "Có", "Không",    "",          "Tối đa 1000 ký tự"],
            ["12", "Nút Lưu",        "Button",       "N/A","N/A",       "",          "Trigger UC 1 — Lưu thông tin KH mới"],
            ["13", "Nút Hủy",        "Button",       "N/A","N/A",       "",          "Hiển thị xác nhận MSG 4 trước khi quay lại"],
        ],
        col_widths_cm=[0.83, 2.8, 2.4, 1.5, 1.8, 1.8, 6.5]
    )

    # ════════════════════════════════════════
    # 5. NON-FUNCTIONAL REQUIREMENTS
    # ════════════════════════════════════════
    apply_heading_style(doc, "5. Yêu cầu phi chức năng (Non-Functional Requirements)", 1)
    add_para(doc, "Phần này mô tả hoạt động của hệ thống về mặt kỹ thuật và hiệu năng.", after=4)

    apply_heading_style(doc, "5.1. Yêu cầu hiệu năng (Performance Requirements)", 2)
    make_table(doc,
        ["Tiêu đề", "Biến số / Tiêu chí", "Ghi chú"],
        [
            ["Thời gian phản hồi", "Trung bình < 2 giây", "95% request phản hồi trong vòng 3 giây"],
            ["Tải trọng",          "200 user đồng thời",   "Ổn định khi 200 user online cùng lúc"],
            ["Khả năng mở rộng",   "Horizontal scaling",   "Hỗ trợ mở rộng server khi user tăng"],
            ["Nền tảng",           "Chrome, Firefox, Edge (latest)", "Responsive trên Desktop và Tablet"],
        ],
        col_widths_cm=[2.21, 6.1, 9.5]
    )

    apply_heading_style(doc, "5.2. Yêu cầu bảo mật (Security Requirements)", 2)
    make_table(doc,
        ["Tiêu đề", "Biến số / Tiêu chí", "Ghi chú"],
        [["Xác thực (Authentication)", "Username/password + OTP", "Hỗ trợ Single Sign-On (SSO)"]],
        col_widths_cm=[2.5, 5.8, 9.5]
    )

    # ════════════════════════════════════════
    # 6. OTHER REQUIREMENTS
    # ════════════════════════════════════════
    apply_heading_style(doc, "6. Yêu cầu khác (Other Requirements)", 1)

    apply_heading_style(doc, "6.1. Cấu hình thông báo chung (Common Messages Configuration)", 2)
    make_table(doc,
        ["Loại thông báo", "Ghi chú"],
        [
            ["In-line Error Message",  "Hiển thị chữ đỏ nghiêng ngay bên dưới trường bị lỗi"],
            ["Error Message",          "Pop-Up chỉ chứa nội dung lỗi và nút Đóng"],
            ["Confirmation Message",   "Pop-Up chứa nội dung và 2 nút: Đồng ý / Hủy"],
            ["Informing Message",      "Pop-Up chứa nội dung thông báo và nút OK"],
            ["Standard platform Message", "Thông báo chuẩn của nền tảng (vd: session timeout)"],
        ],
        col_widths_cm=[3.72, 14.08]
    )

    # ════════════════════════════════════════
    # 7. INTEGRATION
    # ════════════════════════════════════════
    apply_heading_style(doc, "7. Tích hợp (Integration)", 1)
    add_para(doc, "Hệ thống CRM tích hợp với:", after=4)
    add_list_para(doc, "Hệ thống Email (SMTP): Gửi thông báo và email tự động.")
    add_list_para(doc, "Hệ thống ERP: Đồng bộ thông tin khách hàng (chi tiết theo tài liệu tích hợp riêng).")
    doc.add_paragraph()

    # ════════════════════════════════════════
    # 8. DATA MIGRATION
    # ════════════════════════════════════════
    apply_heading_style(doc, "8. Di chuyển dữ liệu (Data Migration)", 1)
    add_para(doc, "Yêu cầu chi tiết cho Di chuyển dữ liệu sẽ được mô tả trong tài liệu SRS riêng. Tài liệu SRS chức năng và SRS di chuyển dữ liệu cần được đọc song song.", after=4)

    # ════════════════════════════════════════
    # 9. APPENDICES
    # ════════════════════════════════════════
    apply_heading_style(doc, "9. Phụ lục (Appendices)", 1)

    apply_heading_style(doc, "9.1. Danh sách thông báo (Messages List)", 2)
    add_para(doc, "Để biết mô tả loại thông báo, tham khảo phần Cấu hình thông báo chung.", after=4)
    make_table(doc,
        ["#", "Mã và nội dung thông báo", "Loại"],
        [
            ["1", "MSG 1: Vui lòng nhập thông tin [Tên trường]",              "Error Message"],
            ["2", "MSG 2: Định dạng [Email] không hợp lệ",                    "Error Message"],
            ["3", "MSG 3: Định dạng [Số điện thoại] không hợp lệ",           "Error Message"],
            ["4", "MSG 4: Bạn có chắc chắn muốn hủy? Thông tin sẽ không được lưu.", "Confirmation"],
            ["5", "MSG 5: [Tên trường] không được vượt quá [N] ký tự",        "In-line Error"],
            ["6", "MSG 6: Email này đã tồn tại trong hệ thống",               "Error Message"],
            ["7", "MSG 7: Bạn có chắc chắn muốn xóa khách hàng này?",        "Confirmation"],
            ["8", "MSG 8: Mã số thuế không đúng định dạng (10 hoặc 13 số)",   "Error Message"],
            ["9", "MSG 9: Tạo thông tin khách hàng thành công",               "Informing Message"],
        ],
        col_widths_cm=[1.31, 12.96, 3.54]
    )

    apply_heading_style(doc, "9.2. Mẫu Email (Email Templates)", 2)
    add_para(doc, "Nội dung email nên viết ở thì quá khứ để trang trọng hơn.", after=4)
    add_para(doc, "ET 1: Gửi email thông báo cho Quản lý khi NV kinh doanh tạo khách hàng mới.", after=4)

    et_table = doc.add_table(rows=4, cols=2)
    et_table.style = 'Table Grid'
    et_data = [
        ("Gửi đến", "Quản lý (Manager)"),
        ("CC",       ""),
        ("Tiêu đề",  "[CRM] Khách hàng mới đã được tạo: <<Tên khách hàng>>"),
        ("Nội dung",
         "Kính gửi <<Tên Quản lý>>,\n\n"
         "Một khách hàng mới đã được tạo trong hệ thống CRM:\n"
         "- Mã khách hàng: <<Mã KH>>\n"
         "- Họ và tên: <<Họ tên>>\n"
         "- Loại khách hàng: <<Loại KH>>\n"
         "- Người tạo: <<Tên nhân viên>>\n"
         "- Ngày tạo: <<Ngày tạo>>\n\n"
         "Vui lòng click vào đây để xem chi tiết.\n\n"
         "Trân trọng,\nHệ thống CRM"),
    ]
    for i, (label, value) in enumerate(et_data):
        c0, c1 = et_table.rows[i].cells[0], et_table.rows[i].cells[1]
        c0.text = ''
        add_run(c0.paragraphs[0], label)
        set_cell_shading(c0, LIGHT_BG)
        set_cell_margins(c0)
        c0.width = Cm(2.04)
        c1.text = ''
        add_run(c1.paragraphs[0], value)
        set_cell_margins(c1)
        c1.width = Cm(15.76)
    set_table_borders(et_table)
    doc.add_paragraph()

    # Note
    p = add_para(doc, "", after=4)
    add_run(p, "Ghi chú: Link ")
    add_run(p, "đây", bold=True, underline=True)
    add_run(p, " dẫn đến màn hình Chi tiết khách hàng tương ứng.")

    return doc


if __name__ == "__main__":
    output = r"c:\Users\Admin\Documents\TestAI\Doc2MD template\SRS_CRM_Tao_Thong_Tin_Khach_Hang_v1.0.docx"
    doc = build_document()
    doc.save(output)
    print(f"✅ Đã tạo file: {output}")
